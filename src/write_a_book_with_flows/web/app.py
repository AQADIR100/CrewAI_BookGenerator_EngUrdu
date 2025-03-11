import streamlit as st
from docx import Document  # This will work after installing python-docx
import os
from write_a_book_with_flows.main import BookFlow, BookState
import subprocess  # Import subprocess to run external commands

def create_docx_from_markdown(content, output_filename):
    try:
        doc = Document()
        lines = content.split('\n')
        
        for line in lines:
            if line.startswith('##'):
                # Add heading for H1 level
                doc.add_heading(line[2:], level=1)
            elif line.startswith('#'):
                # Add heading for H1 level
                doc.add_heading(line[1:], level=1)
            elif line.startswith('**'):
                # Add heading for H1 level
                doc.add_heading(line[1:], level=1)    
            else:
                # Add regular paragraph
                doc.add_paragraph(line)
        
        doc.save(output_filename)
    except Exception as e:
        st.error(f"Error creating DOCX: {str(e)}")
        # Create a simple version with error message
        doc = Document()
        doc.add_paragraph("Error creating full document. Please check the preview below.")
        doc.save(output_filename)

def run_streamlit_app():
    subprocess.run(["streamlit", "run", "book_generator.py"])  # Run the Streamlit command

def main():
    st.title("Book Generator")
    
    st.write("""
    Welcome to the Book Generator! Enter your desired topic and goal, 
    in URDU or ENGLISH,and we'll generate a comprehensive book for you.
    """)
    
    topic = st.text_area("Enter your book topic:", 
                        "")
    
    goal = st.text_area("Enter your book goal:", 
                        "")
    
    if st.button("Generate Book"):
        with st.spinner("Generating your book... This may take a few minutes."):
            try:
                BookFlow.initial_state = BookState(
                    topic=topic,
                    goal=goal,
                    title=topic
                )
                
                book_flow = BookFlow()
                book_flow.kickoff()
                
                with open("book.md", "r", encoding="utf-8") as file:
                    book_content = file.read()
                
                # Create Word document
                docx_filename = "book.docx"
                create_docx_from_markdown(book_content, docx_filename)
                
                st.success("Book generated successfully!")
                
                with open(docx_filename, "rb") as docx_file:
                    docx_bytes = docx_file.read()
                    st.download_button(
                        label="Download Book as Word",
                        data=docx_bytes,
                        file_name=docx_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                
                st.markdown("### Book Preview:")
                st.markdown(book_content)
                
            except Exception as e:
                st.error(f"An error occurred: {str(e)}")

if __name__ == "__main__":
    main()

def run_streamlit_app():
    subprocess.run(["streamlit", "run", "src/write_a_book_with_flows/web/app.py"])

def kickoff():
    obj = run_streamlit_app()
    obj.kickoff()                                                 